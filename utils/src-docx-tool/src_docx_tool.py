import os
import argparse
import json
from pathlib import Path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pronoun_changer import gender_flip
from typing import Callable

ENCODING = 'shift_jis'
CONFIG_FILE = 'config.json'

# Ensure custom_function has function type that takes a string and returns a string
custom_function: Callable[[str], str] = gender_flip

def combine_src_to_docx(output_filename):
    doc = Document()
    for file in Path('.').glob('*.src'):
        doc.add_paragraph(file.name).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph('-' * 40)

        with file.open('r', encoding=ENCODING, errors='ignore') as f:
            for line in f:
                doc.add_paragraph(line)

    doc.save(output_filename)
    print(f"Combined .src files into {output_filename}")

def split_docx_to_src(input_filename):
    # Load config
    if not Path(CONFIG_FILE).exists():
        print(f"Config file '{CONFIG_FILE}' not found.")
        return

    with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
        config = json.load(f)

    script_folder = Path(config.get('script_folder', '.'))
    if not script_folder.exists():
        print(f"Script folder '{script_folder}' not found.")
        return

    doc = Document(input_filename)
    src_files = {}
    current_filename = None
    content_lines = []

    # Parse .docx into a mapping of filename -> list of lines
    for para in doc.paragraphs:
        if para.text.endswith('.src') and len(para.text.strip()) > 4:
            if current_filename and content_lines:  # add prev file content
                src_files[current_filename] = content_lines
            current_filename = para.text.strip()
            content_lines = []
        elif para.text == '-' * 40:
            continue
        else:
            content_lines.append(para.text)

    if current_filename and content_lines:   # add last file content
        src_files[current_filename] = content_lines

    for filename, docx_lines in src_files.items():
        original_path = script_folder / filename
        if not original_path.exists():
            print(f"Original script not found: {original_path}")
            continue

        with original_path.open('r', encoding=ENCODING, errors='ignore') as f:
            original_lines = f.readlines()

        # Merge docx_lines with original_lines:
        merged_lines = []

        # Warning if line counts differ
        if len(original_lines) != len(docx_lines):
            print(f"Warning: Line count mismatch for {filename}. Original: {len(original_lines)}, Docx: {len(docx_lines)}")
        
        for i, line in enumerate(docx_lines):
            if i < len(original_lines):
                original_line = original_lines[i]
                if original_line.lstrip().startswith((';', '#')):
                    merged_lines.append(original_line)
                else:
                    merged_lines.append(line)
            else:
                # if docx has more lines than original
                merged_lines.append(line + '\n')

        # Warning if line counts differ
        if len(merged_lines) != len(docx_lines):
            print(f"Warning: Line count mismatch for {filename}. Docx: {len(docx_lines)}, New file: {len(merged_lines)}")

        with open(filename, 'w', encoding=ENCODING, errors='ignore') as f:
            f.writelines(merged_lines)

        print(f"Recreated {filename} with comment lines restored from original.")

def process_src_files():
    """Process all .src files in the current directory with custom_function, overwriting them."""
    for file in Path('.').glob('*.src'):
        with file.open('r', encoding=ENCODING, errors='ignore') as f:
            lines = f.readlines()

        processed_lines = []
        for line in lines:
            stripped = line.strip()
            if stripped == '' or stripped.startswith(';') or stripped.startswith('#'):
                processed_lines.append(line)
            else:
                processed_line = custom_function(line.rstrip('\n')) + '\n'
                processed_lines.append(processed_line)

        # Warning if line counts differ
        if len(processed_lines) != len(lines):
            print(f"Warning: Line count mismatch for {file.name}. Original: {len(lines)}, Processed: {len(processed_lines)}")

        with file.open('w', encoding=ENCODING, errors='ignore') as f:
            f.writelines(processed_lines)

        print(f"Overwritten: {file.name}")


def main():
    parser = argparse.ArgumentParser(description="Combine or split .src files and .docx documents.")
    subparsers = parser.add_subparsers(dest='command')

    combine_parser = subparsers.add_parser('combine', help='Combine all .src files into a .docx')
    combine_parser.add_argument('output', help='Output .docx filename')

    split_parser = subparsers.add_parser('split', help='Split .docx into .src files with line replacements')
    split_parser.add_argument('input', help='Input .docx filename')

    process_parser = subparsers.add_parser('process', help='Process .src files in-place using custom_function on non-comment, non-blank lines')

    args = parser.parse_args()

    if args.command == 'combine':
        combine_src_to_docx(args.output)
    elif args.command == 'split':
        split_docx_to_src(args.input)
    elif args.command == 'process':
        process_src_files()
    else:
        parser.print_help()

if __name__ == '__main__':
    main()
